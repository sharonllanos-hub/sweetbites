#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generar documentos DOCX para SweetBites
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("Iniciando generación de documentos DOCX...")

# ========================================
# DOCUMENTO 1: REQUERIMIENTOS FUNCIONALES
# ========================================
print("\n1/3 Generando: Requerimientos Funcionales...")

doc1 = Document()

# Título
title = doc1.add_heading('REQUERIMIENTOS FUNCIONALES', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc1.add_paragraph('Sistema de Gestión de Recetas de Postres - SweetBites')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc1.add_paragraph()

# Info
doc1.add_paragraph('Versión: 1.0')
doc1.add_paragraph('Fecha: 11 de Junio 2026')
doc1.add_paragraph('Estado: Implementados y Funcionales')
doc1.add_paragraph()

# Resumen
doc1.add_heading('RESUMEN EJECUTIVO', 1)
doc1.add_paragraph('Este documento lista los 40 requerimientos funcionales implementados en SweetBites.')
doc1.add_paragraph()

# Tabla resumen
table = doc1.add_table(rows=9, cols=4)
table.style = 'Light Grid Accent 1'

headers = table.rows[0].cells
headers[0].text = 'Módulo'
headers[1].text = 'Total RF'
headers[2].text = 'Implementados'
headers[3].text = 'Porcentaje'

data = [
    ('Autenticación y Usuarios', '8', '8', '100%'),
    ('Gestión de Recetas', '12', '12', '100%'),
    ('Favoritos y Colecciones', '4', '4', '100%'),
    ('Valoraciones y Comentarios', '3', '3', '100%'),
    ('Panel de Administración', '8', '8', '100%'),
    ('Categorías', '3', '3', '100%'),
    ('Notificaciones', '2', '2', '100%'),
    ('TOTAL', '40', '40', '100%')
]

for i, row_data in enumerate(data, 1):
    cells = table.rows[i].cells
    for j, value in enumerate(row_data):
        cells[j].text = value

doc1.add_page_break()

# MÓDULO 1: AUTENTICACIÓN
doc1.add_heading('MÓDULO 1: AUTENTICACIÓN Y USUARIOS', 1)

reqs_auth = [
    ('RF-AU-01', 'Registro de Usuario', 'Alta',
     'El sistema permite que nuevos usuarios se registren.',
     ['Email único', 'Contraseña encriptada (bcrypt)', 'Validación de datos', 'Rol automático: usuario'],
     'POST /api/auth/register'),

    ('RF-AU-02', 'Inicio de Sesión', 'Alta',
     'Los usuarios pueden iniciar sesión con email y contraseña.',
     ['Token JWT (7 días)', 'Redirección según rol', 'Mensajes de error claros'],
     'POST /api/auth/login'),

    ('RF-AU-03', 'Cerrar Sesión', 'Media',
     'Los usuarios pueden cerrar su sesión.',
     ['Limpieza de tokens', 'Redirección a inicio'],
     'Frontend: authService.logout()'),

    ('RF-AU-04', 'Ver Perfil', 'Alta',
     'Ver información personal y estadísticas.',
     ['Datos personales', 'Estadísticas (recetas, favoritos)', 'Badges de rol y plan'],
     'GET /api/users/profile'),

    ('RF-AU-05', 'Editar Perfil', 'Alta',
     'Actualizar información personal.',
     ['Editar nombre, teléfono, bio', 'Validación de campos', 'Actualización inmediata'],
     'PUT /api/auth/profile'),

    ('RF-AU-06', 'Cambiar Foto de Perfil', 'Media',
     'Subir o cambiar foto de perfil.',
     ['Formatos: JPG, PNG, WebP', 'Máx 5MB', 'Preview antes de guardar'],
     'POST /api/users/profile/photo'),

    ('RF-AU-07', 'Cambiar Contraseña', 'Alta',
     'Cambiar contraseña de acceso.',
     ['Verificar contraseña actual', 'Validar nueva (mín 6 caracteres)', 'Encriptación bcrypt'],
     'PUT /api/users/change-password'),

    ('RF-AU-08', 'Plan Premium', 'Baja',
     'Actualizar a plan premium.',
     ['Cambio gratis → premium', 'Badge premium', 'Acceso a funciones premium'],
     'POST /api/users/upgrade-premium')
]

for req_id, nombre, prio, desc, criterios, endpoint in reqs_auth:
    doc1.add_heading(f'{req_id}: {nombre}', 2)
    p = doc1.add_paragraph()
    p.add_run('Prioridad: ').bold = True
    p.add_run(f'{prio} | ')
    p.add_run('Estado: ').bold = True
    p.add_run('✓ Implementado')

    p = doc1.add_paragraph()
    p.add_run('Descripción: ').bold = True
    p.add_run(desc)

    doc1.add_paragraph('Criterios de Aceptación:').runs[0].bold = True
    for criterio in criterios:
        doc1.add_paragraph(f'✓ {criterio}', style='List Bullet')

    p = doc1.add_paragraph()
    p.add_run('Endpoint: ').bold = True
    p.add_run(endpoint)

doc1.add_page_break()

# MÓDULO 2: RECETAS
doc1.add_heading('MÓDULO 2: GESTIÓN DE RECETAS', 1)

reqs_recetas = [
    ('RF-RE-01', 'Ver Listado de Recetas', 'Alta',
     'Catálogo de recetas disponibles.',
     ['Grid responsive', 'Paginación (12 por página)', 'Datos: foto, nombre, categoría, dificultad, rating'],
     'GET /api/recipes'),

    ('RF-RE-02', 'Buscar Recetas', 'Alta',
     'Buscar por nombre o ingredientes.',
     ['Búsqueda en tiempo real', 'Por nombre o ingredientes', 'Resultados destacados'],
     'GET /api/recipes?search=query'),

    ('RF-RE-03', 'Filtrar por Categoría', 'Alta',
     'Filtrar recetas por categoría.',
     ['Categorías: Tortas, Galletas, Postres Fríos, Chocolates, Tartas, Cupcakes', 'Conteo por categoría'],
     'GET /api/recipes?categoria=Tortas'),

    ('RF-RE-04', 'Filtrar por Dificultad', 'Media',
     'Filtrar por nivel de dificultad.',
     ['Opciones: Fácil, Intermedio, Difícil', 'Combinable con otros filtros'],
     'GET /api/recipes?dificultad=Facil'),

    ('RF-RE-05', 'Ver Detalle de Receta', 'Alta',
     'Ver detalle completo de una receta.',
     ['Info completa', 'Ingredientes con cantidades', 'Pasos numerados', 'Autor y fecha', 'Calificación', 'Comentarios'],
     'GET /api/recipes/:id'),

    ('RF-RE-06', 'Crear Receta (Wizard 4 pasos)', 'Alta',
     'Crear recetas mediante wizard multi-paso.',
     ['Paso 1: Info básica', 'Paso 2: Ingredientes', 'Paso 3: Pasos', 'Paso 4: Foto', 'Validación por paso'],
     'POST /api/recipes'),

    ('RF-RE-07', 'Editar Receta Propia', 'Alta',
     'Editar recetas propias.',
     ['Solo el autor puede editar', 'Wizard igual a creación', 'Pre-carga de datos'],
     'PUT /api/recipes/:id'),

    ('RF-RE-08', 'Eliminar Receta Propia', 'Media',
     'Eliminar recetas propias.',
     ['Solo el autor', 'Modal de confirmación', 'Eliminación en cascada'],
     'DELETE /api/recipes/:id'),

    ('RF-RE-09', 'Ver Mis Recetas', 'Alta',
     'Ver todas las recetas creadas.',
     ['Estados: publicada, borrador, archivada', 'Acciones: editar, eliminar, ver'],
     'GET /api/recipes/my-recipes'),

    ('RF-RE-10', 'Imprimir Receta', 'Baja',
     'Imprimir en formato optimizado.',
     ['Vista limpia', 'Optimizada para A4', 'Sin navegación'],
     'Frontend: window.print()'),

    ('RF-RE-11', 'Recetas Destacadas', 'Alta',
     'Mostrar recetas destacadas en home.',
     ['Carrusel', 'Ordenadas por rating', 'Máx 6 recetas', 'Auto-scroll'],
     'GET /api/recipes/featured'),

    ('RF-RE-12', 'Recetas Recientes', 'Media',
     'Mostrar recetas recientes.',
     ['Ordenadas por fecha DESC', 'Últimas 12', 'Grid responsive'],
     'GET /api/recipes?sort=recent&limit=12')
]

for req_id, nombre, prio, desc, criterios, endpoint in reqs_recetas:
    doc1.add_heading(f'{req_id}: {nombre}', 2)
    p = doc1.add_paragraph()
    p.add_run('Prioridad: ').bold = True
    p.add_run(f'{prio} | ')
    p.add_run('Estado: ').bold = True
    p.add_run('✓ Implementado')

    p = doc1.add_paragraph()
    p.add_run('Descripción: ').bold = True
    p.add_run(desc)

    doc1.add_paragraph('Criterios:').runs[0].bold = True
    for criterio in criterios:
        doc1.add_paragraph(f'✓ {criterio}', style='List Bullet')

    p = doc1.add_paragraph()
    p.add_run('Endpoint: ').bold = True
    p.add_run(endpoint)

doc1.add_page_break()

# MÓDULO 3: FAVORITOS
doc1.add_heading('MÓDULO 3: FAVORITOS Y COLECCIONES', 1)

reqs_fav = [
    ('RF-FA-01', 'Agregar a Favoritos', 'Alta',
     'Marcar recetas como favoritas.',
     ['Botón corazón', 'Toggle on/off', 'Feedback visual', 'Usuario autenticado'],
     'POST /api/users/favorites/:recipeId'),

    ('RF-FA-02', 'Ver Favoritos', 'Alta',
     'Ver recetas favoritas.',
     ['Ordenadas por fecha DESC', 'Grid responsive', 'Quitar de favoritos'],
     'GET /api/users/favorites'),

    ('RF-FA-03', 'Crear Colección', 'Media',
     'Crear colecciones personalizadas.',
     ['Nombre obligatorio', 'Descripción opcional', 'Nombre único'],
     'POST /api/users/collections'),

    ('RF-FA-04', 'Agregar a Colección', 'Media',
     'Agregar recetas a colecciones.',
     ['Selector de colección', 'Múltiples colecciones', 'Sin duplicados'],
     'POST /api/users/collections/:id/recipes/:recipeId')
]

for req_id, nombre, prio, desc, criterios, endpoint in reqs_fav:
    doc1.add_heading(f'{req_id}: {nombre}', 2)
    p = doc1.add_paragraph()
    p.add_run('Prioridad: ').bold = True
    p.add_run(f'{prio} | ')
    p.add_run('Estado: ').bold = True
    p.add_run('✓ Implementado')

    p = doc1.add_paragraph()
    p.add_run('Descripción: ').bold = True
    p.add_run(desc)

    doc1.add_paragraph('Criterios:').runs[0].bold = True
    for criterio in criterios:
        doc1.add_paragraph(f'✓ {criterio}', style='List Bullet')

    p = doc1.add_paragraph()
    p.add_run('Endpoint: ').bold = True
    p.add_run(endpoint)

doc1.add_page_break()

# MÓDULO 4: VALORACIONES
doc1.add_heading('MÓDULO 4: VALORACIONES Y COMENTARIOS', 1)

reqs_val = [
    ('RF-VC-01', 'Valorar Receta', 'Alta',
     'Calificar con estrellas (1-5).',
     ['Sistema 5 estrellas', 'Una valoración por usuario', 'Promedio en tiempo real'],
     'POST /api/recipes/:id/rate'),

    ('RF-VC-02', 'Comentar Receta', 'Alta',
     'Dejar comentarios.',
     ['Mín 10, máx 500 caracteres', 'Mostrar autor y fecha', 'Ordenados por recientes'],
     'POST /api/recipes/:id/comments'),

    ('RF-VC-03', 'Eliminar Comentario Propio', 'Media',
     'Eliminar comentarios propios.',
     ['Solo el autor', 'Modal confirmación', 'Actualización inmediata'],
     'DELETE /api/comments/:id')
]

for req_id, nombre, prio, desc, criterios, endpoint in reqs_val:
    doc1.add_heading(f'{req_id}: {nombre}', 2)
    p = doc1.add_paragraph()
    p.add_run('Prioridad: ').bold = True
    p.add_run(f'{prio} | ')
    p.add_run('Estado: ').bold = True
    p.add_run('✓ Implementado')

    p = doc1.add_paragraph()
    p.add_run('Descripción: ').bold = True
    p.add_run(desc)

    doc1.add_paragraph('Criterios:').runs[0].bold = True
    for criterio in criterios:
        doc1.add_paragraph(f'✓ {criterio}', style='List Bullet')

    p = doc1.add_paragraph()
    p.add_run('Endpoint: ').bold = True
    p.add_run(endpoint)

doc1.add_page_break()

# MÓDULO 5: ADMIN
doc1.add_heading('MÓDULO 5: PANEL DE ADMINISTRACIÓN', 1)

reqs_admin = [
    ('RF-AD-01', 'Dashboard Admin', 'Alta',
     'Panel con estadísticas.',
     ['Métricas: usuarios, recetas, comentarios', 'Acceso admin', 'Responsive'],
     'GET /api/admin/stats'),

    ('RF-AD-02', 'Gestión de Usuarios', 'Alta',
     'Gestionar usuarios.',
     ['Listado completo', 'Búsqueda', 'Paginación', 'Cambiar rol', 'Eliminar'],
     'GET /api/admin/users'),

    ('RF-AD-03', 'Cambiar Rol', 'Alta',
     'Cambiar rol de usuarios.',
     ['Dropdown usuario/admin', 'Actualización inmediata', 'No puede cambiar propio rol'],
     'PUT /api/admin/users/:id/role'),

    ('RF-AD-04', 'Eliminar Usuario', 'Alta',
     'Eliminar usuarios.',
     ['Modal confirmación', 'Eliminación cascada', 'No auto-eliminación'],
     'DELETE /api/admin/users/:id'),

    ('RF-AD-05', 'Gestión de Recetas', 'Alta',
     'Gestionar todas las recetas.',
     ['Listado completo', 'Editar cualquiera', 'Eliminar', 'Cambiar estado'],
     'GET /api/admin/recipes'),

    ('RF-AD-06', 'Moderación Comentarios', 'Media',
     'Moderar comentarios.',
     ['Listado todos', 'Eliminar inapropiados', 'Confirmación'],
     'GET /api/admin/comments'),

    ('RF-AD-07', 'Gestión Categorías', 'Media',
     'Gestionar categorías.',
     ['Crear, editar, eliminar', 'Nombre único', 'Conteo recetas'],
     'GET /api/admin/categories'),

    ('RF-AD-08', 'Botón Retroceso Admin', 'Baja',
     'Botón volver en páginas admin.',
     ['BackButton en todas', 'Navegación a /admin', 'Responsive'],
     'Frontend: useNavigate(-1)')
]

for req_id, nombre, prio, desc, criterios, endpoint in reqs_admin:
    doc1.add_heading(f'{req_id}: {nombre}', 2)
    p = doc1.add_paragraph()
    p.add_run('Prioridad: ').bold = True
    p.add_run(f'{prio} | ')
    p.add_run('Estado: ').bold = True
    p.add_run('✓ Implementado')

    p = doc1.add_paragraph()
    p.add_run('Descripción: ').bold = True
    p.add_run(desc)

    doc1.add_paragraph('Criterios:').runs[0].bold = True
    for criterio in criterios:
        doc1.add_paragraph(f'✓ {criterio}', style='List Bullet')

    p = doc1.add_paragraph()
    p.add_run('Endpoint: ').bold = True
    p.add_run(endpoint)

# Guardar documento 1
path1 = 'C:/xampp/htdocs/ProSweetBites/appnueva/SweetBites_Requerimientos_Funcionales.docx'
doc1.save(path1)
print(f"[OK] Generado: {path1}")

# ========================================
# DOCUMENTO 2: DIAGRAMA MER
# ========================================
print("\n2/3 Generando: Diagrama MER...")

doc2 = Document()

title = doc2.add_heading('DIAGRAMA MODELO ENTIDAD-RELACIÓN (MER)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc2.add_paragraph('Base de Datos SweetBites')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc2.add_paragraph()
doc2.add_paragraph('Versión: 1.0')
doc2.add_paragraph('Fecha: 11 de Junio 2026')
doc2.add_paragraph()

doc2.add_heading('DESCRIPCIÓN GENERAL', 1)
doc2.add_paragraph('La base de datos sweetbites_db consta de 12 tablas relacionales que gestionan usuarios, recetas, favoritos, colecciones, valoraciones, comentarios, categorías y notificaciones.')

doc2.add_page_break()

# Tablas
doc2.add_heading('ENTIDADES Y ATRIBUTOS', 1)

tablas = [
    {
        'nombre': 'users',
        'descripcion': 'Almacena información de usuarios registrados',
        'campos': [
            ('id', 'INT AUTO_INCREMENT', 'PK'),
            ('nombre', 'VARCHAR(100)', 'NOT NULL'),
            ('email', 'VARCHAR(100)', 'UNIQUE, NOT NULL'),
            ('password_hash', 'VARCHAR(255)', 'NOT NULL'),
            ('telefono', 'VARCHAR(20)', 'NULL'),
            ('rol', "ENUM('usuario','admin')", 'DEFAULT usuario'),
            ('foto_perfil', 'VARCHAR(255)', 'NULL'),
            ('bio', 'TEXT', 'NULL'),
            ('plan', "ENUM('gratis','premium')", 'DEFAULT gratis'),
            ('fecha_registro', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP')
        ]
    },
    {
        'nombre': 'categories',
        'descripcion': 'Categorías de recetas',
        'campos': [
            ('id', 'INT AUTO_INCREMENT', 'PK'),
            ('nombre', 'VARCHAR(100)', 'UNIQUE, NOT NULL'),
            ('icono', 'VARCHAR(50)', 'NULL'),
            ('color', 'VARCHAR(20)', 'NULL'),
            ('descripcion', 'TEXT', 'NULL')
        ]
    },
    {
        'nombre': 'recipes',
        'descripcion': 'Recetas de postres',
        'campos': [
            ('id', 'INT AUTO_INCREMENT', 'PK'),
            ('nombre', 'VARCHAR(200)', 'NOT NULL'),
            ('descripcion', 'TEXT', 'NULL'),
            ('categoria_id', 'INT', 'FK → categories(id)'),
            ('dificultad', "ENUM('Fácil','Intermedio','Difícil')", 'DEFAULT Intermedio'),
            ('tiempo_preparacion', 'INT', 'minutos'),
            ('porciones', 'INT', 'DEFAULT 1'),
            ('foto_principal', 'VARCHAR(255)', 'NULL'),
            ('autor_id', 'INT', 'FK → users(id) ON DELETE SET NULL'),
            ('estado', "ENUM('publicada','borrador','archivada')", 'DEFAULT publicada'),
            ('fecha_creacion', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP')
        ]
    },
    {
        'nombre': 'ingredients',
        'descripcion': 'Ingredientes de cada receta',
        'campos': [
            ('id', 'INT AUTO_INCREMENT', 'PK'),
            ('receta_id', 'INT', 'FK → recipes(id) ON DELETE CASCADE'),
            ('nombre', 'VARCHAR(100)', 'NOT NULL'),
            ('cantidad', 'DECIMAL(10,2)', 'NULL'),
            ('unidad', 'VARCHAR(50)', 'NULL'),
            ('seccion', 'VARCHAR(100)', 'NULL')
        ]
    },
    {
        'nombre': 'steps',
        'descripcion': 'Pasos de preparación',
        'campos': [
            ('id', 'INT AUTO_INCREMENT', 'PK'),
            ('receta_id', 'INT', 'FK → recipes(id) ON DELETE CASCADE'),
            ('numero_paso', 'INT', 'NOT NULL'),
            ('descripcion', 'TEXT', 'NOT NULL'),
            ('foto', 'VARCHAR(255)', 'NULL')
        ]
    },
    {
        'nombre': 'favorites',
        'descripcion': 'Recetas favoritas de usuarios',
        'campos': [
            ('id', 'INT AUTO_INCREMENT', 'PK'),
            ('usuario_id', 'INT', 'FK → users(id) ON DELETE CASCADE'),
            ('receta_id', 'INT', 'FK → recipes(id) ON DELETE CASCADE'),
            ('fecha_guardado', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'),
            ('UNIQUE', '(usuario_id, receta_id)', 'Constraint')
        ]
    },
    {
        'nombre': 'collections',
        'descripcion': 'Colecciones personalizadas de recetas',
        'campos': [
            ('id', 'INT AUTO_INCREMENT', 'PK'),
            ('usuario_id', 'INT', 'FK → users(id) ON DELETE CASCADE'),
            ('nombre', 'VARCHAR(100)', 'NOT NULL'),
            ('descripcion', 'TEXT', 'NULL'),
            ('fecha_creacion', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP')
        ]
    },
    {
        'nombre': 'collection_recipes',
        'descripcion': 'Relación N:M entre colecciones y recetas',
        'campos': [
            ('id', 'INT AUTO_INCREMENT', 'PK'),
            ('coleccion_id', 'INT', 'FK → collections(id) ON DELETE CASCADE'),
            ('receta_id', 'INT', 'FK → recipes(id) ON DELETE CASCADE'),
            ('fecha_agregado', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP')
        ]
    },
    {
        'nombre': 'ratings',
        'descripcion': 'Valoraciones de recetas (1-5 estrellas)',
        'campos': [
            ('id', 'INT AUTO_INCREMENT', 'PK'),
            ('receta_id', 'INT', 'FK → recipes(id) ON DELETE CASCADE'),
            ('usuario_id', 'INT', 'FK → users(id) ON DELETE CASCADE'),
            ('puntuacion', 'INT', 'CHECK (1-5)'),
            ('fecha', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'),
            ('UNIQUE', '(receta_id, usuario_id)', 'Constraint')
        ]
    },
    {
        'nombre': 'comments',
        'descripcion': 'Comentarios en recetas',
        'campos': [
            ('id', 'INT AUTO_INCREMENT', 'PK'),
            ('receta_id', 'INT', 'FK → recipes(id) ON DELETE CASCADE'),
            ('usuario_id', 'INT', 'FK → users(id) ON DELETE CASCADE'),
            ('comentario', 'TEXT', 'NOT NULL'),
            ('fecha', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP')
        ]
    },
    {
        'nombre': 'notifications',
        'descripcion': 'Notificaciones de usuarios',
        'campos': [
            ('id', 'INT AUTO_INCREMENT', 'PK'),
            ('usuario_id', 'INT', 'FK → users(id) ON DELETE CASCADE'),
            ('tipo', 'VARCHAR(50)', 'NOT NULL'),
            ('titulo', 'VARCHAR(200)', 'NOT NULL'),
            ('mensaje', 'TEXT', 'NOT NULL'),
            ('leida', 'BOOLEAN', 'DEFAULT FALSE'),
            ('fecha', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP')
        ]
    }
]

for tabla in tablas:
    doc2.add_heading(f"TABLA: {tabla['nombre']}", 2)
    doc2.add_paragraph(tabla['descripcion'])

    # Tabla de campos
    t = doc2.add_table(rows=len(tabla['campos'])+1, cols=3)
    t.style = 'Light Grid Accent 1'

    headers = t.rows[0].cells
    headers[0].text = 'Campo'
    headers[1].text = 'Tipo'
    headers[2].text = 'Restricciones'

    for i, (campo, tipo, restriccion) in enumerate(tabla['campos'], 1):
        cells = t.rows[i].cells
        cells[0].text = campo
        cells[1].text = tipo
        cells[2].text = restriccion

    doc2.add_paragraph()

doc2.add_page_break()

# Relaciones
doc2.add_heading('RELACIONES ENTRE ENTIDADES', 1)

relaciones = [
    ('users', 'recipes', '1:N', 'Un usuario puede crear muchas recetas (autor)'),
    ('categories', 'recipes', '1:N', 'Una categoría tiene muchas recetas'),
    ('recipes', 'ingredients', '1:N', 'Una receta tiene muchos ingredientes'),
    ('recipes', 'steps', '1:N', 'Una receta tiene muchos pasos'),
    ('users', 'favorites', '1:N', 'Un usuario puede tener muchos favoritos'),
    ('recipes', 'favorites', '1:N', 'Una receta puede estar en muchos favoritos'),
    ('users', 'collections', '1:N', 'Un usuario puede crear muchas colecciones'),
    ('collections', 'collection_recipes', '1:N', 'Una colección tiene muchas recetas'),
    ('recipes', 'collection_recipes', '1:N', 'Una receta puede estar en muchas colecciones'),
    ('users', 'ratings', '1:N', 'Un usuario puede valorar muchas recetas'),
    ('recipes', 'ratings', '1:N', 'Una receta puede tener muchas valoraciones'),
    ('users', 'comments', '1:N', 'Un usuario puede hacer muchos comentarios'),
    ('recipes', 'comments', '1:N', 'Una receta puede tener muchos comentarios'),
    ('users', 'notifications', '1:N', 'Un usuario puede tener muchas notificaciones')
]

t_rel = doc2.add_table(rows=len(relaciones)+1, cols=4)
t_rel.style = 'Light Grid Accent 1'

headers = t_rel.rows[0].cells
headers[0].text = 'Entidad Origen'
headers[1].text = 'Entidad Destino'
headers[2].text = 'Cardinalidad'
headers[3].text = 'Descripción'

for i, (origen, destino, card, desc) in enumerate(relaciones, 1):
    cells = t_rel.rows[i].cells
    cells[0].text = origen
    cells[1].text = destino
    cells[2].text = card
    cells[3].text = desc

# Guardar documento 2
path2 = 'C:/xampp/htdocs/ProSweetBites/appnueva/SweetBites_Diagrama_MER.docx'
doc2.save(path2)
print(f"[OK] Generado: {path2}")

# ========================================
# DOCUMENTO 3: CASOS DE USO
# ========================================
print("\n3/3 Generando: Casos de Uso...")

doc3 = Document()

title = doc3.add_heading('CASOS DE USO', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc3.add_paragraph('Sistema SweetBites')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc3.add_paragraph()
doc3.add_paragraph('Versión: 1.0')
doc3.add_paragraph('Fecha: 11 de Junio 2026')
doc3.add_paragraph()

doc3.add_heading('MÓDULO 1: AUTENTICACIÓN Y USUARIOS', 1)

casos_uso = [
    {
        'id': 'CU-AU-01',
        'nombre': 'Registrarse en el Sistema',
        'actor': 'Usuario no registrado',
        'precondiciones': 'El usuario no tiene cuenta',
        'flujo': [
            'Usuario accede a /auth/register',
            'Sistema muestra formulario de registro',
            'Usuario ingresa: nombre, email, contraseña, teléfono (opcional)',
            'Usuario hace clic en "Registrarse"',
            'Sistema valida datos',
            'Sistema encripta contraseña con bcrypt',
            'Sistema guarda usuario en BD con rol "usuario"',
            'Sistema muestra mensaje de éxito',
            'Sistema redirige a página de login'
        ],
        'postcondiciones': 'Usuario creado en base de datos'
    },
    {
        'id': 'CU-AU-02',
        'nombre': 'Iniciar Sesión',
        'actor': 'Usuario registrado',
        'precondiciones': 'El usuario tiene cuenta',
        'flujo': [
            'Usuario accede a /auth/login',
            'Sistema muestra formulario de login',
            'Usuario ingresa email y contraseña',
            'Usuario hace clic en "Iniciar Sesión"',
            'Sistema valida credenciales',
            'Sistema genera token JWT (válido 7 días)',
            'Sistema almacena token en localStorage',
            'Sistema redirige según rol (admin → /admin, usuario → /)'
        ],
        'postcondiciones': 'Usuario autenticado con sesión activa'
    },
    {
        'id': 'CU-RE-01',
        'nombre': 'Buscar Recetas',
        'actor': 'Usuario (autenticado o no)',
        'precondiciones': 'Existen recetas en el sistema',
        'flujo': [
            'Usuario accede a /recipes',
            'Sistema muestra catálogo de recetas',
            'Usuario ingresa término en barra de búsqueda',
            'Sistema busca en tiempo real (debounce 300ms)',
            'Sistema busca en: nombre de receta, ingredientes',
            'Sistema muestra resultados destacados',
            'Usuario puede hacer clic en una receta para ver detalle'
        ],
        'postcondiciones': 'Usuario ve recetas que coinciden con búsqueda'
    },
    {
        'id': 'CU-RE-02',
        'nombre': 'Crear Receta con Wizard',
        'actor': 'Usuario autenticado',
        'precondiciones': 'Usuario tiene sesión activa',
        'flujo': [
            'Usuario accede a /user/create-recipe',
            'Sistema muestra Paso 1/4: Información Básica',
            'Usuario ingresa: nombre, descripción, categoría, dificultad, tiempo, porciones',
            'Usuario hace clic en "Siguiente"',
            'Sistema valida y muestra Paso 2/4: Ingredientes',
            'Usuario agrega ingredientes con cantidad y unidad',
            'Usuario hace clic en "Siguiente"',
            'Sistema muestra Paso 3/4: Pasos de Preparación',
            'Usuario agrega pasos numerados',
            'Usuario hace clic en "Siguiente"',
            'Sistema muestra Paso 4/4: Foto Principal',
            'Usuario sube foto',
            'Sistema muestra preview',
            'Usuario hace clic en "Publicar"',
            'Sistema guarda receta en BD',
            'Sistema muestra mensaje de éxito'
        ],
        'postcondiciones': 'Receta creada y publicada'
    },
    {
        'id': 'CU-FA-01',
        'nombre': 'Agregar Receta a Favoritos',
        'actor': 'Usuario autenticado',
        'precondiciones': 'Usuario autenticado, receta existe',
        'flujo': [
            'Usuario ve listado de recetas',
            'Usuario hace clic en ícono de corazón en una receta',
            'Sistema valida sesión activa',
            'Sistema agrega receta a favoritos del usuario',
            'Sistema muestra feedback visual (corazón lleno)',
            'Sistema guarda en tabla favorites'
        ],
        'postcondiciones': 'Receta agregada a favoritos del usuario'
    },
    {
        'id': 'CU-VC-01',
        'nombre': 'Valorar Receta',
        'actor': 'Usuario autenticado',
        'precondiciones': 'Usuario autenticado, receta existe',
        'flujo': [
            'Usuario accede a detalle de receta',
            'Sistema muestra sistema de 5 estrellas',
            'Usuario selecciona cantidad de estrellas (1-5)',
            'Sistema valida una valoración por usuario',
            'Sistema guarda valoración en tabla ratings',
            'Sistema recalcula promedio de receta',
            'Sistema actualiza visualización en tiempo real'
        ],
        'postcondiciones': 'Valoración guardada, promedio actualizado'
    },
    {
        'id': 'CU-AD-01',
        'nombre': 'Gestionar Usuarios (Admin)',
        'actor': 'Administrador',
        'precondiciones': 'Usuario con rol admin',
        'flujo': [
            'Admin accede a /admin/users',
            'Sistema muestra listado de usuarios',
            'Sistema muestra: nombre, email, rol, fecha registro',
            'Admin puede buscar por nombre o email',
            'Admin puede cambiar rol (dropdown usuario/admin)',
            'Sistema actualiza rol en BD',
            'Admin puede eliminar usuario',
            'Sistema muestra modal de confirmación',
            'Sistema elimina usuario y datos relacionados (cascada)'
        ],
        'postcondiciones': 'Usuario modificado o eliminado'
    }
]

for cu in casos_uso:
    doc3.add_heading(f"{cu['id']}: {cu['nombre']}", 2)

    p = doc3.add_paragraph()
    p.add_run('Actor: ').bold = True
    p.add_run(cu['actor'])

    p = doc3.add_paragraph()
    p.add_run('Precondiciones: ').bold = True
    p.add_run(cu['precondiciones'])

    doc3.add_paragraph('Flujo Principal:').runs[0].bold = True
    for i, paso in enumerate(cu['flujo'], 1):
        doc3.add_paragraph(f'{i}. {paso}', style='List Number')

    p = doc3.add_paragraph()
    p.add_run('Postcondiciones: ').bold = True
    p.add_run(cu['postcondiciones'])

    doc3.add_paragraph()

# Guardar documento 3
path3 = 'C:/xampp/htdocs/ProSweetBites/appnueva/SweetBites_Casos_de_Uso.docx'
doc3.save(path3)
print(f"[OK] Generado: {path3}")

print("\n" + "="*50)
print("[OK] PROCESO COMPLETADO!")
print("="*50)
print(f"\nDocumentos generados en:")
print(f"  1. {path1}")
print(f"  2. {path2}")
print(f"  3. {path3}")
print("\nFormato: DOCX profesional, limpio, sin colores.")
