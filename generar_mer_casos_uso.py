#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generar documentos de Diagrama MER y Casos de Uso
Formato limpio, Arial 12, sin colores
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_rgb):
    """Establecer color de fondo de celda"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_rgb)
    cell._element.get_or_add_tcPr().append(shading_elm)

print("Generando documentos de Diagrama MER y Casos de Uso...")

# ========================================
# DOCUMENTO 1: DIAGRAMA MER
# ========================================
print("\n1/2 Generando: Diagrama MER...")

doc_mer = Document()

# Configurar márgenes
sections = doc_mer.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Título principal
title = doc_mer.add_heading('DIAGRAMA MODELO ENTIDAD-RELACIÓN (MER)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True

subtitle = doc_mer.add_paragraph('Base de Datos SweetBites')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(12)

doc_mer.add_paragraph()

# Descripción general
p = doc_mer.add_paragraph()
run = p.add_run('La base de datos sweetbites_db está compuesta por 12 entidades relacionales que gestionan usuarios, recetas, favoritos, colecciones, valoraciones, comentarios, categorías y notificaciones.')
run.font.name = 'Arial'
run.font.size = Pt(12)

doc_mer.add_paragraph()
doc_mer.add_page_break()

# Lista de entidades
entidades = [
    {
        'nombre': 'users',
        'descripcion': 'Almacena información de usuarios registrados en el sistema',
        'atributos': [
            ('id', 'INT', 'Clave primaria, autoincremental'),
            ('nombre', 'VARCHAR(100)', 'Nombre completo del usuario'),
            ('email', 'VARCHAR(100)', 'Correo electrónico único'),
            ('password_hash', 'VARCHAR(255)', 'Contraseña encriptada con bcrypt'),
            ('telefono', 'VARCHAR(20)', 'Número de teléfono (opcional)'),
            ('rol', "ENUM('usuario','admin')", 'Rol del usuario'),
            ('foto_perfil', 'VARCHAR(255)', 'URL de la foto de perfil'),
            ('bio', 'TEXT', 'Biografía del usuario'),
            ('plan', "ENUM('gratis','premium')", 'Plan de suscripción'),
            ('fecha_registro', 'DATETIME', 'Fecha y hora de registro')
        ]
    },
    {
        'nombre': 'categories',
        'descripcion': 'Categorías de recetas disponibles en el sistema',
        'atributos': [
            ('id', 'INT', 'Clave primaria, autoincremental'),
            ('nombre', 'VARCHAR(100)', 'Nombre de la categoría (único)'),
            ('icono', 'VARCHAR(50)', 'Icono representativo'),
            ('color', 'VARCHAR(20)', 'Color para la UI'),
            ('descripcion', 'TEXT', 'Descripción de la categoría')
        ]
    },
    {
        'nombre': 'recipes',
        'descripcion': 'Recetas de postres publicadas por usuarios',
        'atributos': [
            ('id', 'INT', 'Clave primaria, autoincremental'),
            ('nombre', 'VARCHAR(200)', 'Nombre de la receta'),
            ('descripcion', 'TEXT', 'Descripción detallada'),
            ('categoria_id', 'INT', 'FK → categories(id)'),
            ('dificultad', "ENUM('Fácil','Intermedio','Difícil')", 'Nivel de dificultad'),
            ('tiempo_preparacion', 'INT', 'Tiempo en minutos'),
            ('porciones', 'INT', 'Número de porciones'),
            ('foto_principal', 'VARCHAR(255)', 'URL de la foto principal'),
            ('autor_id', 'INT', 'FK → users(id) ON DELETE SET NULL'),
            ('estado', "ENUM('publicada','borrador','archivada')", 'Estado de la receta'),
            ('fecha_creacion', 'DATETIME', 'Fecha de creación')
        ]
    },
    {
        'nombre': 'ingredients',
        'descripcion': 'Ingredientes requeridos para cada receta',
        'atributos': [
            ('id', 'INT', 'Clave primaria, autoincremental'),
            ('receta_id', 'INT', 'FK → recipes(id) ON DELETE CASCADE'),
            ('nombre', 'VARCHAR(100)', 'Nombre del ingrediente'),
            ('cantidad', 'DECIMAL(10,2)', 'Cantidad necesaria'),
            ('unidad', 'VARCHAR(50)', 'Unidad de medida (gr, ml, taza, etc.)'),
            ('seccion', 'VARCHAR(100)', 'Sección del ingrediente (opcional)')
        ]
    },
    {
        'nombre': 'steps',
        'descripcion': 'Pasos de preparación de cada receta',
        'atributos': [
            ('id', 'INT', 'Clave primaria, autoincremental'),
            ('receta_id', 'INT', 'FK → recipes(id) ON DELETE CASCADE'),
            ('numero_paso', 'INT', 'Orden del paso'),
            ('descripcion', 'TEXT', 'Descripción del paso'),
            ('foto', 'VARCHAR(255)', 'URL de foto del paso (opcional)')
        ]
    },
    {
        'nombre': 'favorites',
        'descripcion': 'Recetas marcadas como favoritas por usuarios',
        'atributos': [
            ('id', 'INT', 'Clave primaria, autoincremental'),
            ('usuario_id', 'INT', 'FK → users(id) ON DELETE CASCADE'),
            ('receta_id', 'INT', 'FK → recipes(id) ON DELETE CASCADE'),
            ('fecha_guardado', 'DATETIME', 'Fecha en que se guardó'),
            ('UNIQUE', '(usuario_id, receta_id)', 'Un usuario no puede guardar la misma receta dos veces')
        ]
    },
    {
        'nombre': 'collections',
        'descripcion': 'Colecciones personalizadas de recetas creadas por usuarios',
        'atributos': [
            ('id', 'INT', 'Clave primaria, autoincremental'),
            ('usuario_id', 'INT', 'FK → users(id) ON DELETE CASCADE'),
            ('nombre', 'VARCHAR(100)', 'Nombre de la colección'),
            ('descripcion', 'TEXT', 'Descripción (opcional)'),
            ('fecha_creacion', 'DATETIME', 'Fecha de creación')
        ]
    },
    {
        'nombre': 'collection_recipes',
        'descripcion': 'Tabla intermedia que relaciona colecciones con recetas (N:M)',
        'atributos': [
            ('id', 'INT', 'Clave primaria, autoincremental'),
            ('coleccion_id', 'INT', 'FK → collections(id) ON DELETE CASCADE'),
            ('receta_id', 'INT', 'FK → recipes(id) ON DELETE CASCADE'),
            ('fecha_agregado', 'DATETIME', 'Fecha en que se agregó a la colección')
        ]
    },
    {
        'nombre': 'ratings',
        'descripcion': 'Valoraciones de recetas (sistema de 1 a 5 estrellas)',
        'atributos': [
            ('id', 'INT', 'Clave primaria, autoincremental'),
            ('receta_id', 'INT', 'FK → recipes(id) ON DELETE CASCADE'),
            ('usuario_id', 'INT', 'FK → users(id) ON DELETE CASCADE'),
            ('puntuacion', 'INT', 'Valoración del 1 al 5'),
            ('fecha', 'DATETIME', 'Fecha de valoración'),
            ('UNIQUE', '(receta_id, usuario_id)', 'Un usuario solo puede valorar una receta una vez')
        ]
    },
    {
        'nombre': 'comments',
        'descripcion': 'Comentarios de usuarios en recetas',
        'atributos': [
            ('id', 'INT', 'Clave primaria, autoincremental'),
            ('receta_id', 'INT', 'FK → recipes(id) ON DELETE CASCADE'),
            ('usuario_id', 'INT', 'FK → users(id) ON DELETE CASCADE'),
            ('comentario', 'TEXT', 'Texto del comentario'),
            ('fecha', 'DATETIME', 'Fecha del comentario')
        ]
    },
    {
        'nombre': 'notifications',
        'descripcion': 'Notificaciones del sistema para usuarios',
        'atributos': [
            ('id', 'INT', 'Clave primaria, autoincremental'),
            ('usuario_id', 'INT', 'FK → users(id) ON DELETE CASCADE'),
            ('tipo', 'VARCHAR(50)', 'Tipo de notificación'),
            ('titulo', 'VARCHAR(200)', 'Título de la notificación'),
            ('mensaje', 'TEXT', 'Mensaje completo'),
            ('leida', 'BOOLEAN', 'Estado de lectura'),
            ('fecha', 'DATETIME', 'Fecha de creación')
        ]
    }
]

# Generar tabla para cada entidad
heading = doc_mer.add_heading('ENTIDADES DE LA BASE DE DATOS', 1)
for run in heading.runs:
    run.font.name = 'Arial'

for entidad in entidades:
    # Nombre de la entidad
    h = doc_mer.add_heading(f"ENTIDAD: {entidad['nombre']}", 2)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(14)

    # Descripción
    p = doc_mer.add_paragraph()
    run = p.add_run(entidad['descripcion'])
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    # Tabla de atributos
    table = doc_mer.add_table(rows=len(entidad['atributos'])+1, cols=3)
    table.style = 'Table Grid'

    # Encabezados
    headers = table.rows[0].cells
    for cell in headers:
        set_cell_background(cell, 'E8E8E8')

    p = headers[0].paragraphs[0]
    run = p.add_run('Atributo')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    p = headers[1].paragraphs[0]
    run = p.add_run('Tipo de Dato')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    p = headers[2].paragraphs[0]
    run = p.add_run('Descripción')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    # Datos
    for i, (nombre, tipo, desc) in enumerate(entidad['atributos'], 1):
        cells = table.rows[i].cells

        p = cells[0].paragraphs[0]
        run = p.add_run(nombre)
        run.font.name = 'Arial'
        run.font.size = Pt(11)

        p = cells[1].paragraphs[0]
        run = p.add_run(tipo)
        run.font.name = 'Arial'
        run.font.size = Pt(11)

        p = cells[2].paragraphs[0]
        run = p.add_run(desc)
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    doc_mer.add_paragraph()

doc_mer.add_page_break()

# Relaciones entre entidades
heading = doc_mer.add_heading('RELACIONES ENTRE ENTIDADES', 1)
for run in heading.runs:
    run.font.name = 'Arial'

relaciones = [
    ('users', 'recipes', '1:N', 'Un usuario puede crear muchas recetas (como autor)'),
    ('categories', 'recipes', '1:N', 'Una categoría contiene muchas recetas'),
    ('recipes', 'ingredients', '1:N', 'Una receta tiene muchos ingredientes'),
    ('recipes', 'steps', '1:N', 'Una receta tiene muchos pasos de preparación'),
    ('users', 'favorites', '1:N', 'Un usuario puede tener muchas recetas favoritas'),
    ('recipes', 'favorites', '1:N', 'Una receta puede ser favorita de muchos usuarios'),
    ('users', 'collections', '1:N', 'Un usuario puede crear muchas colecciones'),
    ('collections', 'collection_recipes', '1:N', 'Una colección contiene muchas recetas'),
    ('recipes', 'collection_recipes', '1:N', 'Una receta puede estar en muchas colecciones'),
    ('users', 'ratings', '1:N', 'Un usuario puede valorar muchas recetas'),
    ('recipes', 'ratings', '1:N', 'Una receta puede tener muchas valoraciones'),
    ('users', 'comments', '1:N', 'Un usuario puede hacer muchos comentarios'),
    ('recipes', 'comments', '1:N', 'Una receta puede tener muchos comentarios'),
    ('users', 'notifications', '1:N', 'Un usuario puede tener muchas notificaciones')
]

# Tabla de relaciones
table = doc_mer.add_table(rows=len(relaciones)+1, cols=4)
table.style = 'Table Grid'

# Encabezados
headers = table.rows[0].cells
for cell in headers:
    set_cell_background(cell, 'E8E8E8')

p = headers[0].paragraphs[0]
run = p.add_run('Entidad Origen')
run.font.name = 'Arial'
run.font.size = Pt(12)
run.font.bold = True

p = headers[1].paragraphs[0]
run = p.add_run('Entidad Destino')
run.font.name = 'Arial'
run.font.size = Pt(12)
run.font.bold = True

p = headers[2].paragraphs[0]
run = p.add_run('Cardinalidad')
run.font.name = 'Arial'
run.font.size = Pt(12)
run.font.bold = True

p = headers[3].paragraphs[0]
run = p.add_run('Descripción')
run.font.name = 'Arial'
run.font.size = Pt(12)
run.font.bold = True

# Datos
for i, (origen, destino, card, desc) in enumerate(relaciones, 1):
    cells = table.rows[i].cells

    p = cells[0].paragraphs[0]
    run = p.add_run(origen)
    run.font.name = 'Arial'
    run.font.size = Pt(11)

    p = cells[1].paragraphs[0]
    run = p.add_run(destino)
    run.font.name = 'Arial'
    run.font.size = Pt(11)

    p = cells[2].paragraphs[0]
    run = p.add_run(card)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True

    p = cells[3].paragraphs[0]
    run = p.add_run(desc)
    run.font.name = 'Arial'
    run.font.size = Pt(11)

# Guardar documento MER
path_mer = 'C:/xampp/htdocs/ProSweetBites/appnueva/SweetBites_Diagrama_MER_Completo.docx'
doc_mer.save(path_mer)
print(f"[OK] Generado: {path_mer}")

# ========================================
# DOCUMENTO 2: CASOS DE USO
# ========================================
print("\n2/2 Generando: Casos de Uso...")

doc_cu = Document()

# Configurar márgenes
sections = doc_cu.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Título principal
title = doc_cu.add_heading('CASOS DE USO', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True

subtitle = doc_cu.add_paragraph('Sistema SweetBites - Gestión de Recetas de Postres')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(12)

doc_cu.add_paragraph()

# Casos de uso
casos_uso = [
    {
        'id': 'CU-AU-01',
        'nombre': 'Registrarse en el Sistema',
        'modulo': 'Autenticación y Usuarios',
        'actor': 'Usuario no registrado',
        'precondiciones': 'El usuario no tiene una cuenta en el sistema',
        'flujo': [
            'El usuario accede a la página de registro (/auth/register)',
            'El sistema muestra el formulario de registro',
            'El usuario ingresa su nombre completo',
            'El usuario ingresa su correo electrónico',
            'El usuario crea una contraseña segura',
            'El usuario confirma la contraseña',
            'El usuario opcionalmente ingresa su teléfono',
            'El usuario hace clic en "Registrarse"',
            'El sistema valida que todos los campos sean correctos',
            'El sistema verifica que el correo no esté ya registrado',
            'El sistema encripta la contraseña con bcrypt',
            'El sistema crea el usuario en la base de datos con rol "usuario"',
            'El sistema muestra mensaje de éxito',
            'El sistema redirige al usuario a la página de inicio de sesión'
        ],
        'postcondiciones': 'El usuario queda registrado en el sistema y puede iniciar sesión'
    },
    {
        'id': 'CU-AU-02',
        'nombre': 'Iniciar Sesión',
        'modulo': 'Autenticación y Usuarios',
        'actor': 'Usuario registrado',
        'precondiciones': 'El usuario tiene una cuenta registrada en el sistema',
        'flujo': [
            'El usuario accede a la página de login (/auth/login)',
            'El sistema muestra el formulario de inicio de sesión',
            'El usuario ingresa su correo electrónico',
            'El usuario ingresa su contraseña',
            'El usuario hace clic en "Iniciar Sesión"',
            'El sistema valida las credenciales contra la base de datos',
            'El sistema genera un token JWT válido por 7 días',
            'El sistema almacena el token en localStorage',
            'Si el usuario es admin, redirige a /admin',
            'Si el usuario es normal, redirige a la página principal'
        ],
        'postcondiciones': 'El usuario queda autenticado con sesión activa'
    },
    {
        'id': 'CU-RE-01',
        'nombre': 'Buscar Recetas',
        'modulo': 'Gestión de Recetas',
        'actor': 'Usuario (autenticado o no)',
        'precondiciones': 'Existen recetas publicadas en el sistema',
        'flujo': [
            'El usuario accede al catálogo de recetas (/recipes)',
            'El sistema muestra el listado de recetas',
            'El usuario escribe un término en la barra de búsqueda',
            'El sistema espera 300ms (debounce) antes de buscar',
            'El sistema busca en nombres de recetas e ingredientes',
            'El sistema muestra los resultados que coinciden',
            'Los resultados se destacan visualmente',
            'El usuario puede hacer clic en una receta para ver su detalle'
        ],
        'postcondiciones': 'El usuario visualiza las recetas que coinciden con su búsqueda'
    },
    {
        'id': 'CU-RE-02',
        'nombre': 'Crear Receta con Wizard',
        'modulo': 'Gestión de Recetas',
        'actor': 'Usuario autenticado',
        'precondiciones': 'El usuario tiene sesión activa',
        'flujo': [
            'El usuario accede a /user/create-recipe',
            'El sistema muestra el Paso 1/4: Información Básica',
            'El usuario ingresa nombre, descripción, categoría, dificultad, tiempo y porciones',
            'El usuario hace clic en "Siguiente"',
            'El sistema valida los datos y muestra el Paso 2/4: Ingredientes',
            'El usuario agrega ingredientes con cantidad y unidad',
            'El usuario hace clic en "Siguiente"',
            'El sistema muestra el Paso 3/4: Pasos de Preparación',
            'El usuario agrega los pasos numerados',
            'El usuario hace clic en "Siguiente"',
            'El sistema muestra el Paso 4/4: Foto Principal',
            'El usuario sube una foto',
            'El sistema muestra preview de la foto',
            'El usuario hace clic en "Publicar"',
            'El sistema guarda la receta en la base de datos',
            'El sistema muestra mensaje de éxito',
            'El sistema redirige al detalle de la receta creada'
        ],
        'postcondiciones': 'La receta queda creada y publicada en el sistema'
    },
    {
        'id': 'CU-FA-01',
        'nombre': 'Agregar Receta a Favoritos',
        'modulo': 'Favoritos y Colecciones',
        'actor': 'Usuario autenticado',
        'precondiciones': 'El usuario está autenticado y existe la receta',
        'flujo': [
            'El usuario visualiza el listado o detalle de recetas',
            'El usuario identifica el ícono de corazón en una receta',
            'El usuario hace clic en el ícono de corazón',
            'El sistema valida que el usuario esté autenticado',
            'El sistema agrega la receta a la tabla favorites',
            'El sistema muestra feedback visual (corazón lleno)',
            'El sistema muestra mensaje de confirmación'
        ],
        'postcondiciones': 'La receta queda guardada en favoritos del usuario'
    },
    {
        'id': 'CU-FA-02',
        'nombre': 'Ver Mis Favoritos',
        'modulo': 'Favoritos y Colecciones',
        'actor': 'Usuario autenticado',
        'precondiciones': 'El usuario tiene sesión activa',
        'flujo': [
            'El usuario accede a /user/favorites',
            'El sistema obtiene las recetas favoritas del usuario',
            'El sistema las ordena por fecha de guardado descendente',
            'El sistema muestra el listado en grid responsive',
            'El usuario puede hacer clic en una receta para ver su detalle',
            'El usuario puede quitar recetas de favoritos haciendo clic en el corazón'
        ],
        'postcondiciones': 'El usuario visualiza todas sus recetas favoritas'
    },
    {
        'id': 'CU-VC-01',
        'nombre': 'Valorar Receta',
        'modulo': 'Valoraciones y Comentarios',
        'actor': 'Usuario autenticado',
        'precondiciones': 'El usuario está autenticado y existe la receta',
        'flujo': [
            'El usuario accede al detalle de una receta',
            'El sistema muestra el sistema de 5 estrellas',
            'El usuario selecciona la cantidad de estrellas (1 a 5)',
            'El sistema valida que el usuario no haya valorado antes esta receta',
            'El sistema guarda la valoración en la tabla ratings',
            'El sistema recalcula el promedio de valoraciones',
            'El sistema actualiza la visualización en tiempo real',
            'El sistema muestra mensaje de confirmación'
        ],
        'postcondiciones': 'La valoración queda registrada y el promedio actualizado'
    },
    {
        'id': 'CU-VC-02',
        'nombre': 'Comentar Receta',
        'modulo': 'Valoraciones y Comentarios',
        'actor': 'Usuario autenticado',
        'precondiciones': 'El usuario está autenticado y existe la receta',
        'flujo': [
            'El usuario accede al detalle de una receta',
            'El usuario se desplaza a la sección de comentarios',
            'El usuario escribe su comentario en la caja de texto',
            'El sistema valida que tenga entre 10 y 500 caracteres',
            'El usuario hace clic en "Comentar"',
            'El sistema guarda el comentario en la base de datos',
            'El sistema muestra el comentario con nombre, foto y fecha del usuario',
            'Los comentarios se ordenan por más recientes primero'
        ],
        'postcondiciones': 'El comentario queda publicado en la receta'
    },
    {
        'id': 'CU-AD-01',
        'nombre': 'Gestionar Usuarios (Administrador)',
        'modulo': 'Panel de Administración',
        'actor': 'Administrador',
        'precondiciones': 'El usuario tiene rol de administrador',
        'flujo': [
            'El administrador accede a /admin/users',
            'El sistema muestra el listado de todos los usuarios',
            'El sistema muestra nombre, email, rol y fecha de registro',
            'El administrador puede buscar por nombre o email',
            'El administrador puede cambiar el rol usando el dropdown',
            'El sistema actualiza el rol en la base de datos',
            'El administrador puede eliminar un usuario haciendo clic en "Eliminar"',
            'El sistema muestra modal de confirmación con advertencia',
            'El administrador confirma la eliminación',
            'El sistema elimina el usuario y sus datos relacionados (cascada)',
            'El sistema muestra mensaje de confirmación'
        ],
        'postcondiciones': 'Los usuarios quedan gestionados según las acciones del administrador'
    },
    {
        'id': 'CU-AD-02',
        'nombre': 'Moderar Comentarios (Administrador)',
        'modulo': 'Panel de Administración',
        'actor': 'Administrador',
        'precondiciones': 'El usuario tiene rol de administrador',
        'flujo': [
            'El administrador accede a /admin/comments',
            'El sistema muestra el listado de todos los comentarios',
            'El sistema muestra autor, receta asociada, contenido y fecha',
            'Los comentarios están ordenados por más recientes primero',
            'El administrador identifica un comentario inapropiado',
            'El administrador hace clic en "Eliminar"',
            'El sistema muestra modal de confirmación',
            'El administrador confirma la eliminación',
            'El sistema elimina el comentario de la base de datos',
            'El sistema actualiza el listado',
            'El sistema muestra mensaje de confirmación'
        ],
        'postcondiciones': 'Los comentarios inapropiados quedan eliminados del sistema'
    }
]

# Generar tabla para cada caso de uso
for idx, cu in enumerate(casos_uso, 1):
    # Crear tabla de 6 filas x 2 columnas
    table = doc_cu.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    # Configurar ancho de columnas
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)

    # FILA 1: ID y Nombre
    cell_header = table.rows[0].cells[0]
    cell_header.merge(table.rows[0].cells[1])
    set_cell_background(cell_header, 'E8E8E8')
    p = cell_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{cu['id']}  {cu['nombre']}")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    # FILA 2: Módulo
    cell_label = table.rows[1].cells[0]
    set_cell_background(cell_label, 'F5F5F5')
    p = cell_label.paragraphs[0]
    run = p.add_run('Módulo:')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    cell_value = table.rows[1].cells[1]
    p = cell_value.paragraphs[0]
    run = p.add_run(cu['modulo'])
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    # FILA 3: Actor
    cell_label = table.rows[2].cells[0]
    set_cell_background(cell_label, 'F5F5F5')
    p = cell_label.paragraphs[0]
    run = p.add_run('ACTOR')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    cell_value = table.rows[2].cells[1]
    p = cell_value.paragraphs[0]
    run = p.add_run(cu['actor'])
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    # FILA 4: Precondiciones
    cell_label = table.rows[3].cells[0]
    set_cell_background(cell_label, 'F5F5F5')
    p = cell_label.paragraphs[0]
    run = p.add_run('PRECONDICIONES')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    cell_value = table.rows[3].cells[1]
    p = cell_value.paragraphs[0]
    run = p.add_run(cu['precondiciones'])
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    # FILA 5: Flujo Principal
    cell_label = table.rows[4].cells[0]
    set_cell_background(cell_label, 'F5F5F5')
    p = cell_label.paragraphs[0]
    run = p.add_run('FLUJO PRINCIPAL')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    cell_value = table.rows[4].cells[1]
    p = cell_value.paragraphs[0]
    for i, paso in enumerate(cu['flujo'], 1):
        if i > 1:
            p = cell_value.add_paragraph()
        run = p.add_run(f'{i}. {paso}')
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    # FILA 6: Postcondiciones
    cell_label = table.rows[5].cells[0]
    set_cell_background(cell_label, 'F5F5F5')
    p = cell_label.paragraphs[0]
    run = p.add_run('POSTCONDICIONES')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    cell_value = table.rows[5].cells[1]
    p = cell_value.paragraphs[0]
    run = p.add_run(cu['postcondiciones'])
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    # Espacio entre tablas
    doc_cu.add_paragraph()

    # Salto de página cada 2 casos de uso
    if idx % 2 == 0 and idx < len(casos_uso):
        doc_cu.add_page_break()

# Guardar documento Casos de Uso
path_cu = 'C:/xampp/htdocs/ProSweetBites/appnueva/SweetBites_Casos_de_Uso_Completos.docx'
doc_cu.save(path_cu)
print(f"[OK] Generado: {path_cu}")

print("\n" + "="*60)
print("[OK] PROCESO COMPLETADO!")
print("="*60)
print(f"\nDocumentos generados:")
print(f"  1. {path_mer}")
print(f"  2. {path_cu}")
print("\nFormato: Arial 12, sin colores, tablas limpias y profesionales")
